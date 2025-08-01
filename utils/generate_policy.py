from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

def generate_policy_docx(doc: Document, data: dict):
    replace_placeholder_text(doc, "XXXXXXXXXXX", data.get("company", "某保险公司"))
    replace_placeholder_text(doc, "{{PRICE_INFO}}", f"{data.get('total_premium', '$XXX')}/{data.get('policy_term', '6个月')}，一次性付款")

    # 写入责任险
    write_checkbox_and_amount(doc, "Liability", data["liability"]["selected"])
    if data["liability"]["selected"]:
        replace_text_in_paragraphs(doc, "赔偿对方医疗费最高$XXXX/人", f"赔偿对方医疗费最高{data['liability']['bi_per_person']}/人")
        replace_text_in_paragraphs(doc, "赔偿对方医疗费总额最高$XXX", f"赔偿对方医疗费总额最高{data['liability']['bi_per_accident']}")
        replace_text_in_paragraphs(doc, "赔偿对方车辆和财产损失最多$XXXX", f"赔偿对方车辆和财产损失最多{data['liability']['pd']}")
    else:
        replace_text_in_paragraphs(doc, "赔偿对方医疗费最高$XXXX/人", "没有选择该项目")
        replace_text_in_paragraphs(doc, "赔偿对方医疗费总额最高$XXX", "")
        replace_text_in_paragraphs(doc, "赔偿对方车辆和财产损失最多$XXXX", "")

    # 无保险驾驶者
    write_checkbox_and_amount(doc, "Uninsured Motorist", data["uninsured_motorist"]["selected"])
    if data["uninsured_motorist"]["selected"]:
        if data["uninsured_motorist"].get("bi_per_person"):
            replace_text_in_paragraphs(doc, "赔偿你和乘客医疗费$XXXX/人", f"赔偿你和乘客医疗费{data['uninsured_motorist']['bi_per_person']}/人")
        if data["uninsured_motorist"].get("bi_per_accident"):
            replace_text_in_paragraphs(doc, "一场事故最多赔偿医疗费$XXXX", f"一场事故最多赔偿医疗费{data['uninsured_motorist']['bi_per_accident']}")
        if data["uninsured_motorist"].get("pd"):
            replace_text_in_paragraphs(doc, "赔偿自己车辆最多$XXX(自付额$250)", f"赔偿自己车辆最多{data['uninsured_motorist']['pd']}(自付额$250)")
    else:
        replace_text_in_paragraphs(doc, "赔偿你和乘客医疗费$XXXX/人", "没有选择该项目")
        replace_text_in_paragraphs(doc, "一场事故最多赔偿医疗费$XXXX", "")
        replace_text_in_paragraphs(doc, "赔偿自己车辆最多$XXX(自付额$250)", "")

    # Medical Payment
    write_checkbox_and_amount(doc, "Medical Payment", data["medical_payment"]["selected"])
    if data["medical_payment"]["selected"]:
        replace_text_in_paragraphs(doc, "赔偿自己和自己车上乘客在事故中受伤的医疗费每人$XXX",
                                   f"赔偿自己和自己车上乘客在事故中受伤的医疗费每人{data['medical_payment']['med']}")
    else:
        replace_text_in_paragraphs(doc, "赔偿自己和自己车上乘客在事故中受伤的医疗费每人$XXX", "没有选择该项目")

    # Personal Injury
    write_checkbox_and_amount(doc, "Personal Injury", data["personal_injury"]["selected"])
    if data["personal_injury"]["selected"]:
        replace_text_in_paragraphs(doc, "赔偿自己和自己车上乘客在事故中受伤的医疗费，误工费和精神损失费每人$XXX",
                                   f"赔偿自己和自己车上乘客在事故中受伤的医疗费，误工费和精神损失费每人{data['personal_injury']['pip']}")
    else:
        replace_text_in_paragraphs(doc, "赔偿自己和自己车上乘客在事故中受伤的医疗费，误工费和精神损失费每人$XXX", "没有选择该项目")

    insert_vehicle_section(doc, data.get("vehicles", []))

def insert_vehicle_section(doc: Document, vehicles: list):
    if not vehicles:
        return

    marker_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "车辆保障:" in p.text:
            marker_idx = i
            break
    if marker_idx == -1:
        return

    marker_p = doc.paragraphs[marker_idx]
    marker_el = marker_p._element

    next_el = marker_el.getnext()
    while next_el is not None and (next_el.tag.endswith("p") or next_el.tag.endswith("tbl")):
        to_remove = next_el
        next_el = next_el.getnext()
        marker_el.getparent().remove(to_remove)

    vehicle_table_template = None
    for tbl in doc.tables:
        if "Collision" in tbl.cell(1, 0).text and "租车报销" in tbl.cell(4, 0).text:
            vehicle_table_template = tbl
            break
    if not vehicle_table_template:
        return

    for vehicle in vehicles:
        spacer_p = marker_p.insert_paragraph_after("·")
        spacer_p.runs[0].font.size = Pt(1)
        spacer_p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

        vin_text = f"{vehicle['model']}     VIN：{vehicle['vin']}"
        vin_p = spacer_p.insert_paragraph_after(vin_text)
        vin_p.runs[0].font.size = Pt(12)
        vin_p.runs[0].bold = True

        new_table = deepcopy(vehicle_table_template._element)
        vin_p._element.addnext(new_table)
        new_table_obj = vin_p._element.getnext()

        fill_vehicle_table(doc, new_table_obj, vehicle)
        marker_p = doc.paragraphs[-1]

def fill_vehicle_table(doc: Document, table_el, vehicle: dict):
    tbl = None
    for t in doc.tables:
        if t._element == table_el:
            tbl = t
            break
    if not tbl:
        return

    update_checkbox_cell(tbl.cell(1, 1), vehicle["collision"]["selected"])
    update_checkbox_cell(tbl.cell(2, 1), vehicle["comprehensive"]["selected"])
    update_checkbox_cell(tbl.cell(3, 1), vehicle["roadside"]["selected"])
    update_checkbox_cell(tbl.cell(4, 1), vehicle["rental"]["selected"])

    if vehicle["collision"]["selected"]:
        tbl.cell(1, 2).text = f"""自付额${vehicle['collision']['deductible']}
修车时自付额以内自己出，自付额以外的保险公司赔付"""
    else:
        tbl.cell(1, 2).text = "没有选择该项目"

    if vehicle["comprehensive"]["selected"]:
        tbl.cell(2, 2).text = f"""自付额${vehicle['comprehensive']['deductible']}
修车时自付额以内自己出，自付额以外的保险公司赔付"""
    else:
        tbl.cell(2, 2).text = "没有选择该项目"

    if vehicle["roadside"]["selected"]:
        tbl.cell(3, 2).text = "赔偿由于:机械故障,电瓶没电,钥匙被锁车内,燃油耗尽,轮胎没气造成车辆不可行驶时的免费拖车，免费充电，免费开锁服务"
    else:
        tbl.cell(3, 2).text = "没有选择该项目"

    if vehicle["rental"]["selected"]:
        tbl.cell(4, 2).text = "每天$30 最多30天"
    else:
        tbl.cell(4, 2).text = "没有选择该项目"

def update_checkbox_cell(cell, selected):
    cell.text = "✅" if selected else "❌"
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(16)

def replace_placeholder_text(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)

def replace_text_in_paragraphs(doc, old, new):
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)

def write_checkbox_and_amount(doc, keyword, selected):
    symbol = "✅" if selected else "❌"
    for table in doc.tables:
        for row in table.rows:
            if keyword in row.cells[0].text:
                cell = row.cells[1]
                cell.text = symbol
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(16)
